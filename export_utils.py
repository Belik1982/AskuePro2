import pandas as pd
from fpdf import FPDF
import os
import tempfile
import io
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import numpy as np
from docx import Document # Новая библиотека
from docx.shared import Pt, RGBColor

FONT_NAME = "DejaVuSans.ttf"

class PDFReport(FPDF):
    def __init__(self, title_text, file_info=None):
        super().__init__()
        self.title_text = title_text
        self.file_info = file_info
        self.font_loaded = False
        if os.path.exists(FONT_NAME):
            try:
                self.add_font('DejaVu', '', FONT_NAME)
                self.add_font('DejaVu', 'B', FONT_NAME)
                self.font_loaded = True
            except: pass

    def header(self):
        font = 'DejaVu' if self.font_loaded else 'Arial'
        self.set_font(font, 'B', 14)
        self.cell(0, 10, self._txt(self.title_text), ln=True, align='C')
        self.ln(5)

    def footer(self):
        self.set_y(-15)
        font = 'DejaVu' if self.font_loaded else 'Arial'
        self.set_font(font, '', 8)
        self.cell(0, 10, f'Page {self.page_no()}', align='C')

    def _txt(self, text):
        if self.font_loaded: return str(text)
        try: return str(text).encode('latin-1', 'replace').decode('latin-1')
        except: return str(text)

    def add_section_header(self, title):
        font = 'DejaVu' if self.font_loaded else 'Arial'
        self.set_font(font, 'B', 12)
        self.set_fill_color(230, 230, 230)
        self.cell(0, 8, self._txt(title), ln=True, fill=True)
        self.ln(3)

    def add_image_from_file(self, img_path):
        try:
            if self.get_y() > 220: self.add_page()
            self.image(img_path, x=10, w=190)
        except Exception as e:
            self.cell(0, 10, self._txt(f"Img Error: {e}"), ln=True)
        self.ln(5)

    def add_stats_table(self, stats_dict):
        font = 'DejaVu' if self.font_loaded else 'Arial'
        self.set_font(font, '', 10)
        for k, v in stats_dict.items():
            self.cell(70, 6, self._txt(k), border=1)
            self.cell(0, 6, self._txt(str(v)), border=1, ln=True)
        self.ln(5)

# --- ГЕНЕРАТОРЫ ГРАФИКОВ ---
def render_mpl_chart(df, title):
    try:
        fig, ax = plt.subplots(figsize=(10, 4))
        for (meter, typ), group in df.groupby(['MeterID', 'Type']):
            ax.plot(group['DateTime'], group['Value'], label=f"{meter} {typ}", linewidth=1)
        ax.set_title(title)
        ax.set_ylabel("кВт / кВАр")
        ax.grid(True, linestyle='--', alpha=0.5)
        ax.xaxis.set_major_formatter(mdates.DateFormatter('%d.%m %H:%M'))
        plt.setp(ax.get_xticklabels(), rotation=45, ha="right", fontsize=8)
        if len(df.groupby(['MeterID', 'Type'])) < 10:
            ax.legend(loc='upper right', fontsize='x-small', framealpha=0.5)
        plt.tight_layout()
        tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        plt.savefig(tmp.name, dpi=100)
        plt.close(fig)
        return tmp.name
    except: return None

def render_mpl_daily(df, title):
    try:
        daily = df.groupby(["Date", "MeterID", "Type"])["Value"].sum().reset_index()
        daily = daily.sort_values("Date")
        fig, ax = plt.subplots(figsize=(10, 4))
        for (meter, typ), group in daily.groupby(['MeterID', 'Type']):
            ax.bar(group['Date'], group['Value'], label=f"{meter} {typ}", alpha=0.7)
        ax.set_title(title)
        ax.set_ylabel("кВт*ч")
        ax.grid(True, axis='y', linestyle='--', alpha=0.5)
        ax.xaxis.set_major_formatter(mdates.DateFormatter('%d.%m'))
        ax.xaxis.set_major_locator(mdates.DayLocator(interval=1))
        plt.setp(ax.get_xticklabels(), rotation=90, fontsize=8)
        if len(daily.groupby(['MeterID', 'Type'])) < 10:
            ax.legend(loc='upper right', fontsize='x-small')
        plt.tight_layout()
        tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        plt.savefig(tmp.name, dpi=100)
        plt.close(fig)
        return tmp.name
    except: return None

def render_mpl_matrix(df, title):
    try:
        if df.empty: return None
        m_id = df.iloc[0]['MeterID']
        t_id = df.iloc[0]['Type']
        sub = df[(df['MeterID'] == m_id) & (df['Type'] == t_id)].copy()
        sub['Hour'] = sub['DateTime'].dt.hour
        sub['Day'] = sub['DateTime'].dt.date
        piv = sub.pivot_table(index='Hour', columns='Day', values='Value', aggfunc='mean')
        fig, ax = plt.subplots(figsize=(10, 5))
        cax = ax.imshow(piv, aspect='auto', cmap='viridis', interpolation='nearest')
        fig.colorbar(cax, label='кВт')
        ax.set_title(title)
        ax.set_ylabel("Час")
        ax.set_xlabel("День")
        ax.set_yticks(np.arange(0, 24, 2))
        days = piv.columns
        if len(days) > 0:
            step = max(1, len(days) // 15)
            ax.set_xticks(np.arange(0, len(days), step))
            ax.set_xticklabels([d.strftime('%d.%m') for d in days[::step]], rotation=45, fontsize=8)
        if len(days) <= 35:
            rows, cols = piv.shape
            for i in range(rows):
                for j in range(cols):
                    val = piv.iloc[i, j]
                    if not np.isnan(val) and val > 0:
                        ax.text(j, i, f"{val:.0f}", ha="center", va="center", color="white", fontsize=5)
        plt.tight_layout()
        tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        plt.savefig(tmp.name, dpi=120)
        plt.close(fig)
        return tmp.name
    except: return None

def export_custom_pdf(df_full, file_info, config) -> bytes:
    pdf = PDFReport(config.get('title', 'Отчет'), file_info)
    pdf.add_page()
    d_start, d_end = config['dates']
    df_period = df_full[(df_full['Date'] >= d_start) & (df_full['Date'] <= d_end)].copy()
    font = 'DejaVu' if pdf.font_loaded else 'Arial'
    pdf.set_font(font, '', 9)
    pdf.cell(0, 5, pdf._txt(f"Период отчета: {d_start} - {d_end}"), ln=True)
    f_names = ", ".join([f['name'] for f in file_info[:10]])
    if len(file_info) > 10: f_names += "..."
    pdf.multi_cell(0, 5, pdf._txt(f"Файлы: {f_names}"))
    pdf.ln(5)
    
    for i, block in enumerate(config['blocks']):
        b_type = block['type']
        title = block.get('title', f'Блок {i+1}')
        meters = block.get('meters', [])
        types = block.get('types', [])
        if not meters or not types: continue
        mask = df_period['MeterID'].isin(meters) & df_period['Type'].isin(types)
        df_block = df_period[mask]
        if df_block.empty: continue
        
        pdf.add_section_header(f"{i+1}. {title}")
        img_path = None
        if b_type == 'stats':
            pdf.set_font(font, '', 10)
            total = df_block['Value'].sum()
            peak = df_block['Value'].max()
            pdf.cell(60, 6, pdf._txt("Сумма по выборке:"), border=1)
            pdf.cell(60, 6, f"{total:,.0f}".replace(",", " "), border=1, ln=True)
            pdf.cell(60, 6, pdf._txt("Максимум:"), border=1)
            pdf.cell(60, 6, f"{peak:,.0f}".replace(",", " "), border=1, ln=True)
            pdf.ln(5)
        elif b_type == 'graph_30m':
            img_path = render_mpl_chart(df_block, title)
        elif b_type == 'graph_daily':
            img_path = render_mpl_daily(df_block, title)
        elif b_type == 'graph_matrix':
            img_path = render_mpl_matrix(df_block, title)
        if img_path:
            pdf.add_image_from_file(img_path)
            try: os.unlink(img_path)
            except: pass
            
    return bytes(pdf.output(dest='S'))

def export_excel_bytes(df: pd.DataFrame, include_index=False) -> bytes:
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        df.to_excel(writer, index=include_index, sheet_name="Data")
    return buf.getvalue()

# --- EXPORT CHAT TO DOCX ---
def export_chat_to_docx(messages):
    """Экспорт истории чата в Word"""
    doc = Document()
    doc.add_heading('Історія чату з ШІ-асистентом', 0)
    
    # Пропускаем первое сообщение (системный промпт)
    visible_msgs = messages[1:] if len(messages) > 0 else []
    
    for msg in visible_msgs:
        role_name = "ВИ" if msg["role"] == "user" else "ШІ (АСИСТЕНТ)"
        
        # Заголовок сообщения
        p = doc.add_paragraph()
        run = p.add_run(f"{role_name}:")
        run.bold = True
        run.font.size = Pt(11)
        if msg["role"] == "model" or msg["role"] == "assistant":
            run.font.color.rgb = RGBColor(0, 100, 0) # Зеленый для ИИ
        else:
            run.font.color.rgb = RGBColor(0, 0, 150) # Синий для юзера
            
        # Текст сообщения
        doc.add_paragraph(msg["content"])
        doc.add_paragraph("-" * 30) # Разделитель

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()